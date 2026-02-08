import io
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from openai import OpenAI
import anthropic


# =========================
# Page Config
# =========================
st.set_page_config(page_title="Agentic Resume Assistant", layout="centered")


# =========================
# Keys / Clients
# =========================
OPENAI_KEY = st.secrets.get("OPENAI_API_KEY", "")
ANTHROPIC_KEY = st.secrets.get("ANTHROPIC_API_KEY", "")

if not OPENAI_KEY:
    st.error("Missing OPENAI_API_KEY in Streamlit secrets.")
    st.stop()

if not ANTHROPIC_KEY:
    st.error("Missing ANTHROPIC_API_KEY in Streamlit secrets.")
    st.stop()

client_openai = OpenAI(api_key=OPENAI_KEY)
client_claude = anthropic.Anthropic(api_key=ANTHROPIC_KEY)


# =========================
# Session State (persist results across reruns)
# =========================
for k in ["resume_text", "feedback", "updated_doc_bytes"]:
    st.session_state.setdefault(k, None)

def reset_outputs():
    st.session_state["resume_text"] = None
    st.session_state["feedback"] = None
    st.session_state["updated_doc_bytes"] = None


# =========================
# Utility: Bullet cleanup
# =========================
def clean_bullets(text: str) -> list[str]:
    if not text:
        return []

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    bullets: list[str] = []

    for ln in lines:
        ln2 = ln.lstrip("•").lstrip("-").strip()
        if len(ln2) >= 3 and ln2[0].isdigit() and ln2[1:3] in [". ", ") "]:
            ln2 = ln2[3:].strip()
        if ln2:
            bullets.append(ln2)

    return bullets[:3]


# =========================
# OpenAI: Bullet generation
# =========================
def generate_bullet_points(subject: str, description: str, github_url: str) -> list[str]:
    prompt = f"""You are a resume expert. Based on the project below, generate 2–3 strong, concise resume bullet points.
Use action verbs, include concrete scope/tech/metrics when possible, and keep each bullet to one line.

Project Title: {subject}
Project Description: {description}
GitHub (optional): {github_url}

Return ONLY the bullet points, one per line, each starting with "• ".
"""

    model_name = st.secrets.get("OPENAI_MODEL", "gpt-4o-mini")
    resp = client_openai.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
    )
    raw = (resp.choices[0].message.content or "").strip()
    bullets = clean_bullets(raw)

    if len(bullets) < 2:
        bullets = [
            "Built an end-to-end project deliverable from requirements through validation, emphasizing clarity and measurable impact.",
            "Implemented reliable data processing and analysis workflow with clean documentation and reproducibility.",
        ][:3]

    return bullets


# =========================
# DOCX: Replace first project under PROJECT EXPERIENCE
# =========================
def replace_first_project_safely(doc: Document, new_title: str, new_bullets: list[str]) -> Document:
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def format_title(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

    def format_bullet(paragraph, text):
        run = paragraph.add_run(f"• {text}")
        run.font.size = Pt(10.5)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

    new_bullets = [bp.strip() for bp in new_bullets if bp and bp.strip()]

    section_found = False
    start_idx = -1
    end_idx = -1

    for i, para in enumerate(doc.paragraphs):
        if "PROJECT EXPERIENCE" in para.text.upper():
            section_found = True
            continue

        if section_found and start_idx == -1 and para.text.strip():
            start_idx = i
            continue

        if section_found and start_idx != -1:
            if para.runs and para.runs[0].bold:
                end_idx = i
                break

    if not section_found:
        raise ValueError("Could not find 'PROJECT EXPERIENCE' section in the document.")
    if start_idx == -1:
        raise ValueError("Found 'PROJECT EXPERIENCE' but couldn't locate the first project entry below it.")

    if end_idx == -1:
        for j in range(start_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.strip() == "":
                end_idx = j
                break
        else:
            end_idx = len(doc.paragraphs)

    for idx in reversed(range(start_idx, end_idx)):
        delete_paragraph(doc.paragraphs[idx])

    insert_idx = start_idx

    for bullet in reversed(new_bullets):
        bullet_para = doc.paragraphs[insert_idx].insert_paragraph_before("")
        format_bullet(bullet_para, bullet)

    title_para = doc.paragraphs[insert_idx].insert_paragraph_before("")
    format_title(title_para, new_title)

    return doc


def extract_text_from_docx(docx_file) -> str:
    doc = Document(docx_file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


# =========================
# Anthropic: List models + feedback
# =========================
@st.cache_data(ttl=300)
def list_anthropic_models() -> list[str]:
    try:
        page = client_claude.models.list(limit=100)
        return [m.id for m in page.data if getattr(m, "id", None)]
    except Exception:
        return []


def get_resume_feedback_from_claude(resume_text: str, model_id: str) -> str:
    system_prompt = "You're a career coach reviewing resumes for clarity, impact, and relevance."
    user_prompt = f"""Evaluate the following resume:

{resume_text}

Give me:
1. 3–5 specific improvement suggestions
2. Weak or vague bullet points, if any
3. Suggestions for tailoring to roles like: data analyst, product manager, ML engineer.
Return your response in a clear bullet list.
"""

    resp = client_claude.messages.create(
        model=model_id,
        system=system_prompt,
        max_tokens=1000,
        temperature=0.4,
        messages=[{"role": "user", "content": user_prompt}],
    )

    return "".join(
        block.text for block in resp.content
        if getattr(block, "type", None) == "text"
    ).strip()


# =========================
# UI
# =========================
st.title("🤖 Agentic AI Resume Assistant")
st.markdown("Upload your resume, replace the first project, and get OpenAI + Claude feedback.")

uploaded_file = st.file_uploader("📄 Upload your `.docx` resume", type=["docx"], key="resume_uploader")

# ✅ If user clicks X (clears uploader), reset and stop
if uploaded_file is None:
    reset_outputs()
    st.info("Upload a resume to begin.")
    st.stop()

st.success("✅ Resume uploaded successfully!")

st.subheader("🛠️ Replace First Project")
subject = st.text_input(
    "Project Title",
    placeholder="Business Analytics Toolbox – Trends and Transitions in Men's College Basketball | Jan 2024 – May 2024",
)
description = st.text_area("Project Description", height=150)
github_url = st.text_input("GitHub Repository URL (optional)")

st.subheader("🧠 Claude Model")
available_models = list_anthropic_models()

if available_models:
    default_model = st.secrets.get("ANTHROPIC_MODEL", available_models[0])
    if default_model not in available_models:
        default_model = available_models[0]
    claude_model = st.selectbox(
        "Pick a Claude model (this list is what your API key can access):",
        options=available_models,
        index=available_models.index(default_model),
    )
else:
    st.warning(
        "Could not list Anthropic models for this key. "
        "We will try a fallback model ID, but it may fail if your key lacks access."
    )
    claude_model = st.secrets.get("ANTHROPIC_MODEL", "claude-sonnet-4-5")

if st.button("✨ Update Resume & Get Feedback"):
    if not subject.strip():
        st.error("Please enter a Project Title.")
        st.stop()

    with st.spinner("Generating bullet points using OpenAI..."):
        bullet_points = generate_bullet_points(subject, description, github_url)

    with st.spinner("Replacing the first project in your resume..."):
        doc = Document(uploaded_file)
        updated_doc = replace_first_project_safely(doc, subject, bullet_points)

        buf = io.BytesIO()
        updated_doc.save(buf)
        updated_bytes = buf.getvalue()

        resume_text = extract_text_from_docx(io.BytesIO(updated_bytes))

    with st.spinner(f"Getting feedback from Claude ({claude_model})..."):
        try:
            feedback = get_resume_feedback_from_claude(resume_text, claude_model)
        except anthropic.NotFoundError:
            st.error(
                f"Model '{claude_model}' is not available for your API key. "
                "Pick a different model from the dropdown (if available), or check your Anthropic plan/access."
            )
            st.stop()

    # ✅ Persist outputs so they do NOT disappear on download click (rerun)
    st.session_state["updated_doc_bytes"] = updated_bytes
    st.session_state["resume_text"] = resume_text
    st.session_state["feedback"] = feedback

# ✅ STATIC RENDER (always show if present)
if st.session_state["resume_text"]:
    st.subheader("✅ Updated Resume Preview")
    st.text_area("Resume Text", st.session_state["resume_text"], height=300)

    st.download_button(
        label="📥 Download Updated Resume",
        data=st.session_state["updated_doc_bytes"],
        file_name="Updated_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.subheader("💬 Claude's Feedback")
    st.markdown(st.session_state["feedback"])
